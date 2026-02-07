"""
文件管理器：管理文件操作（整理、重命名、分类等）

遵循 docs/ARCHITECTURE.md 中的Executor模块规范
"""

from typing import Any, Dict, List, Optional
import logging
from pathlib import Path
import shutil
from datetime import datetime
import time
from agent.tools.exceptions import FileManagerError
from agent.tools.config import Config

logger = logging.getLogger(__name__)


class FileManager:
    """
    文件管理器：执行文件整理、重命名、分类等操作
    
    职责：
    - 文件重命名
    - 文件分类（按类型、日期等）
    - 文件移动和整理
    - 去重处理
    """
    
    def __init__(self, config: Config):
        """
        初始化文件管理器
        
        Args:
            config: 配置对象
        """
        self.config = config
        self.sandbox_path = Path(config.sandbox_path).resolve()
        self.sandbox_path.mkdir(parents=True, exist_ok=True)
        
        # 安全：定义允许操作的目录白名单
        # 放宽限制：允许用户主目录下的所有路径（但禁止系统关键路径）
        home = Path.home()
        self.allowed_paths = [
            self.sandbox_path,
            home,  # 允许操作用户主目录下的所有路径
            # 保留常用路径的明确列表（向后兼容）
            home / "Desktop",
            home / "Downloads",
            home / "Documents",
        ]
        
        # 安全：禁止操作的路径
        self.forbidden_paths = [
            Path("/"),
            Path("/System"),
            Path("/Library"),
            Path("/Applications"),
            Path("/usr"),
            Path("/bin"),
            Path("/sbin"),
            Path("/etc"),
            Path("/var"),
            Path("/private"),
        ]
        
        logger.info(f"文件管理器已初始化，沙盒目录: {self.sandbox_path}")
    
    def _validate_path(self, file_path: Path) -> Path:
        """
        验证文件路径是否安全
        
        安全策略：
        1. 优先检查是否在用户主目录下（允许）
        2. 检查是否在沙盒目录下（允许）
        3. 检查是否是禁止的系统关键路径
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析后的绝对路径
            
        Raises:
            FileManagerError: 如果路径不安全
        """
        # 转换为绝对路径
        abs_path = file_path.resolve()
        home = Path.home()
        
        # 优先检查：是否在用户主目录下（允许）
        try:
            abs_path.relative_to(home)
            # 在用户主目录下，允许操作
            return abs_path
        except ValueError:
            pass
        
        # 检查是否在沙盒目录下（允许）
        try:
            abs_path.relative_to(self.sandbox_path.resolve())
            return abs_path
        except ValueError:
            pass
        
        # 检查是否是禁止的系统关键路径
        # 注意：先检查允许路径，再检查禁止路径，避免误判
        for forbidden in self.forbidden_paths:
            if forbidden == Path("/"):
                # 特殊处理：根路径 / 不应该匹配所有路径
                # 只检查路径是否直接是根路径的子路径（排除 /Users）
                if abs_path == Path("/"):
                    raise FileManagerError(f"禁止操作系统关键路径: {abs_path}")
                # 如果路径在 /Users 下，已经在上面允许了
                if len(abs_path.parts) > 1 and abs_path.parts[1] == "Users":
                    continue
                # 其他根路径下的直接子目录（如 /System, /Library）是禁止的
                if len(abs_path.parts) == 2:
                    if abs_path.parts[1] in ["System", "Library", "Applications", "usr", "bin", "sbin", "etc", "var", "private"]:
                        raise FileManagerError(f"禁止操作系统关键路径: {abs_path}")
            else:
                # 对于其他禁止路径，检查是否在路径下
                try:
                    abs_path.relative_to(forbidden)
                    raise FileManagerError(f"禁止操作系统关键路径: {abs_path}")
                except ValueError:
                    pass  # 不在禁止路径内，继续检查
        
        # 如果都不匹配，拒绝操作
        raise FileManagerError(
            f"路径不在允许的操作范围内: {abs_path}。"
            f"只允许操作用户主目录（{home}）和沙盒目录（{self.sandbox_path}）下的文件。"
        )
    
    def _safe_move(self, source: Path, target: Path) -> None:
        """
        安全地移动文件，带错误处理
        
        Args:
            source: 源文件路径
            target: 目标文件路径
            
        Raises:
            FileManagerError: 如果移动失败
        """
        try:
            # 确保目标目录存在
            target.parent.mkdir(parents=True, exist_ok=True)
            
            # 验证源文件存在且是文件
            if not source.exists():
                raise FileManagerError(f"源文件不存在: {source}")
            
            if not source.is_file():
                raise FileManagerError(f"源路径不是文件: {source}")
            
            # 使用copy+remove而不是move，更安全
            shutil.copy2(source, target)
            
            # 验证目标文件已成功创建
            if not target.exists():
                raise FileManagerError(f"复制失败，目标文件不存在: {target}")
            
            # 验证文件大小一致
            if source.stat().st_size != target.stat().st_size:
                raise FileManagerError("文件大小不一致，复制可能失败")
            
            # 只有在确认复制成功后才删除源文件
            source.unlink()
            logger.info(f"文件已安全移动: {source} -> {target}")
            
        except Exception as e:
            # 如果出错，确保目标文件被清理（如果存在）
            if target.exists():
                try:
                    target.unlink()
                except Exception:
                    pass
            
            raise FileManagerError(f"移动文件失败: {e}") from e
    
    def execute_step(self, step: Dict[str, Any], context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        执行文件操作步骤
        
        Args:
            step: 任务步骤，包含type、action、params等
            context: 上下文信息（可选，用于传递文件缓存等）
        
        Returns:
            执行结果
        """
        step_type = step.get("type")
        action = step.get("action", "")
        params = step.get("params", {})
        
        logger.info(f"执行文件操作: {step_type} - {action}")
        
        try:
            if step_type == "file_organize":
                return self._organize(params)
            elif step_type == "file_rename":
                return self._rename(params)
            elif step_type == "file_move":
                return self._move(params)
            elif step_type == "file_classify":
                return self._classify(params)
            elif step_type == "file_batch_rename":
                return self._batch_rename(params)
            elif step_type == "file_copy":
                return self._copy(params)
            elif step_type == "file_batch_copy":
                return self._batch_copy(params)
            elif step_type == "file_batch_organize":
                return self._batch_organize_files(params)
            elif step_type == "file_create":
                return self._create_file(params)
            elif step_type == "file_write":
                return self._write_file(params)
            elif step_type == "file_read":
                return self._read_file(params)
            elif step_type == "file_delete":
                return self._delete_file(params)
            else:
                raise FileManagerError(f"未知的文件操作类型: {step_type}")
                
        except Exception as e:
            logger.error(f"执行文件操作失败: {e}", exc_info=True)
            return {
                "success": False,
                "message": f"操作失败: {e}",
                "data": None
            }
    
    def _organize(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        整理文件：按类型或日期分类
        
        Args:
            params: 包含file_path（文件路径）、organize_by（分类方式：type/date）
        
        Returns:
            整理结果
        """
        file_path_str = params.get("file_path", "")
        if not file_path_str or file_path_str == "." or file_path_str == "..":
            raise FileManagerError(f"无效的文件路径: {file_path_str}")
        
        file_path = Path(file_path_str)
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        organize_by = params.get("organize_by", "type")
        
        if not file_path.exists():
            raise FileManagerError(f"文件不存在: {file_path}")
        
        if not file_path.is_file():
            raise FileManagerError(f"路径不是文件: {file_path}")
        
        if organize_by == "type":
            # 按文件类型分类
            ext = file_path.suffix.lower() or "unknown"
            category = ext[1:] if ext.startswith(".") else ext
            target_dir = self.sandbox_path / "organized" / category
        elif organize_by == "date":
            # 按日期分类
            date_str = datetime.now().strftime("%Y-%m-%d")
            target_dir = self.sandbox_path / "organized" / date_str
        else:
            raise FileManagerError(f"未知的分类方式: {organize_by}")
        
        target_dir.mkdir(parents=True, exist_ok=True)
        target_path = target_dir / file_path.name
        
        # 如果目标文件已存在，添加序号
        if target_path.exists():
            counter = 1
            stem = file_path.stem
            suffix = file_path.suffix
            while target_path.exists():
                target_path = target_dir / f"{stem}_{counter}{suffix}"
                counter += 1
        
        # 安全：使用安全的移动方法
        self._safe_move(file_path, target_path)
        
        return {
            "success": True,
            "message": f"文件已整理到: {target_path}",
            "data": {"source": str(file_path), "target": str(target_path)}
        }
    
    def _batch_organize_files(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        批量整理文件：按关键词匹配分类，支持时间过滤
        
        Args:
            params: 包含以下参数：
                - source_dir: 源目录路径（如 ~/Downloads）
                - target_base_dir: 目标基础目录（如 ~/Documents，可选，默认使用source_dir）
                - rules: 分类规则列表，每个规则包含：
                    - keywords: 关键词列表（文件名包含这些词则分类）
                    - category: 分类名称（如"发票"、"合同"）
                    - extensions: 文件扩展名列表（如["pdf", "docx"]）
                - days: 只处理最近N天的文件（可选，None表示处理所有文件）
                - file_extensions: 只处理指定扩展名的文件（可选，如["pdf", "jpg"]）
        
        Returns:
            整理结果，包含：
                - success: 是否成功
                - message: 结果消息
                - data: 包含统计信息和详细报告
        """
        source_dir_str = params.get("source_dir", "")
        if not source_dir_str:
            raise FileManagerError("缺少source_dir参数")
        
        # 处理 ~ 符号和路径解析
        if source_dir_str.startswith("~/"):
            source_dir = Path.home() / source_dir_str[2:]
        elif source_dir_str.startswith("~"):
            source_dir = Path.home() / source_dir_str[1:]
        else:
            source_dir = Path(source_dir_str)
        
        if not source_dir.is_absolute():
            source_dir = Path.home() / source_dir
        
        source_dir = source_dir.resolve()
        source_dir = self._validate_path(source_dir)
        
        if not source_dir.exists():
            raise FileManagerError(f"源目录不存在: {source_dir}")
        
        if not source_dir.is_dir():
            raise FileManagerError(f"源路径不是目录: {source_dir}")
        
        # 目标基础目录（默认使用源目录）
        target_base_dir_str = params.get("target_base_dir", str(source_dir))
        if target_base_dir_str.startswith("~/"):
            target_base_dir = Path.home() / target_base_dir_str[2:]
        elif target_base_dir_str.startswith("~"):
            target_base_dir = Path.home() / target_base_dir_str[1:]
        else:
            target_base_dir = Path(target_base_dir_str)
        
        if not target_base_dir.is_absolute():
            target_base_dir = Path.home() / target_base_dir
        
        target_base_dir = target_base_dir.resolve()
        target_base_dir = self._validate_path(target_base_dir)
        target_base_dir.mkdir(parents=True, exist_ok=True)
        
        # 获取分类规则
        rules = params.get("rules", [])
        if not rules:
            # 默认规则：按文件类型分类
            rules = [
                {
                    "keywords": ["发票", "receipt", "invoice"],
                    "category": "发票",
                    "extensions": ["pdf", "docx", "xlsx", "jpg", "png"]
                },
                {
                    "keywords": ["合同", "contract", "agreement"],
                    "category": "合同",
                    "extensions": ["pdf", "docx", "doc"]
                },
                {
                    "keywords": ["简历", "resume", "cv"],
                    "category": "简历",
                    "extensions": ["pdf", "docx", "doc"]
                }
            ]
        
        # 时间过滤（最近N天）
        days = params.get("days")
        cutoff_time = None
        if days is not None:
            cutoff_time = time.time() - (days * 24 * 60 * 60)
        
        # 文件扩展名过滤
        allowed_extensions = params.get("file_extensions")
        if allowed_extensions:
            allowed_extensions = [ext.lower().lstrip(".") for ext in allowed_extensions]
        
        # 统计信息
        stats = {
            "total_scanned": 0,
            "total_moved": 0,
            "by_category": {},
            "skipped": 0,
            "errors": []
        }
        
        # 详细报告
        report = []
        
        # 扫描源目录
        try:
            for file_path in source_dir.iterdir():
                if not file_path.is_file():
                    continue
                
                stats["total_scanned"] += 1
                
                # 时间过滤
                if cutoff_time:
                    file_mtime = file_path.stat().st_mtime
                    if file_mtime < cutoff_time:
                        stats["skipped"] += 1
                        continue
                
                # 扩展名过滤
                file_ext = file_path.suffix.lower().lstrip(".")
                if allowed_extensions and file_ext not in allowed_extensions:
                    stats["skipped"] += 1
                    continue
                
                # 匹配分类规则
                file_name_lower = file_path.name.lower()
                matched_category = None
                
                for rule in rules:
                    # 检查关键词
                    keywords_match = any(
                        keyword.lower() in file_name_lower 
                        for keyword in rule.get("keywords", [])
                    )
                    
                    # 检查扩展名
                    extensions = rule.get("extensions", [])
                    extensions_normalized = [ext.lower().lstrip(".") for ext in extensions]
                    extension_match = not extensions or file_ext in extensions_normalized
                    
                    if keywords_match and extension_match:
                        matched_category = rule.get("category", "其他")
                        break
                
                # 如果没有匹配到规则，使用默认分类（按扩展名）
                if not matched_category:
                    if file_ext in ["jpg", "jpeg", "png", "gif", "bmp", "webp"]:
                        matched_category = "图片"
                    elif file_ext in ["pdf", "doc", "docx", "txt", "rtf"]:
                        matched_category = "文档"
                    elif file_ext in ["mp4", "avi", "mov", "mkv"]:
                        matched_category = "视频"
                    elif file_ext in ["mp3", "wav", "flac", "aac"]:
                        matched_category = "音频"
                    else:
                        matched_category = "其他"
                
                # 创建目标目录
                target_category_dir = target_base_dir / matched_category
                target_category_dir.mkdir(parents=True, exist_ok=True)
                
                # 移动文件
                target_path = target_category_dir / file_path.name
                
                # 处理文件名冲突
                if target_path.exists():
                    counter = 1
                    stem = file_path.stem
                    suffix = file_path.suffix
                    while target_path.exists():
                        target_path = target_category_dir / f"{stem}_{counter}{suffix}"
                        counter += 1
                
                try:
                    shutil.move(str(file_path), str(target_path))
                    stats["total_moved"] += 1
                    stats["by_category"][matched_category] = stats["by_category"].get(matched_category, 0) + 1
                    
                    report.append({
                        "file": file_path.name,
                        "source": str(file_path),
                        "target": str(target_path),
                        "category": matched_category
                    })
                    
                    logger.info(f"✅ 已移动文件: {file_path.name} → {target_path}")
                    
                except Exception as e:
                    error_msg = f"移动文件失败 {file_path.name}: {e}"
                    logger.error(error_msg)
                    stats["errors"].append(error_msg)
            
            # 生成结果消息
            category_summary = ", ".join([
                f"{cat}({count})" 
                for cat, count in stats["by_category"].items()
            ])
            
            message = f"批量整理完成: 扫描{stats['total_scanned']}个文件，移动{stats['total_moved']}个文件"
            if category_summary:
                message += f"，分类: {category_summary}"
            if stats["skipped"] > 0:
                message += f"，跳过{stats['skipped']}个文件"
            if stats["errors"]:
                message += f"，{len(stats['errors'])}个错误"
            
            return {
                "success": True,
                "message": message,
                "data": {
                    "stats": stats,
                    "report": report,
                    "source_dir": str(source_dir),
                    "target_base_dir": str(target_base_dir)
                }
            }
            
        except Exception as e:
            logger.error(f"批量整理文件失败: {e}", exc_info=True)
            raise FileManagerError(f"批量整理文件失败: {e}") from e
    
    def _find_file(self, filename: str, search_dirs: List[Path] = None) -> Optional[Path]:
        """
        智能搜索文件：在指定目录中搜索文件名
        
        支持：
        - 精确匹配（如 "screenshot.png"）
        - 部分匹配（如 "照片" 可以匹配 "照片.png"）
        - 忽略扩展名（如 "照片" 可以匹配 "照片.png", "照片.jpg" 等）
        
        Args:
            filename: 文件名（如 "screenshot.png" 或 "照片"）
            search_dirs: 搜索目录列表，默认为用户主目录下的常用目录
        
        Returns:
            找到的文件路径，如果未找到返回None
        """
        if search_dirs is None:
            home = Path.home()
            search_dirs = [
                home / "Desktop",
                home / "Downloads",
                home / "Documents",
                self.sandbox_path,
                home,  # 也在主目录下搜索
            ]
        
        # 规范化文件名（去除前后空格）
        filename = filename.strip()
        filename_lower = filename.lower()  # 用于大小写不敏感匹配
        
        # 1. 先尝试直接匹配文件名（精确匹配，大小写敏感）
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            candidate = search_dir / filename
            if candidate.exists() and candidate.is_file():
                logger.info(f"精确匹配找到文件: {candidate}")
                return candidate
        
        # 2. 精确匹配（大小写不敏感）
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            try:
                for item in search_dir.iterdir():
                    if item.is_file() and item.name.lower() == filename_lower:
                        logger.info(f"精确匹配（大小写不敏感）找到文件: {item}")
                        return item
            except (PermissionError, OSError):
                continue
        
        # 3. 如果文件名不包含扩展名，尝试添加常见扩展名（优先尝试）
        if "." not in filename:
            common_extensions = [".docx", ".doc", ".pdf", ".xls", ".xlsx", ".txt", ".png", ".jpg", ".jpeg", ".gif", ".zip", ".rar"]
            for ext in common_extensions:
                for search_dir in search_dirs:
                    if not search_dir.exists():
                        continue
                    candidate = search_dir / f"{filename}{ext}"
                    if candidate.exists() and candidate.is_file():
                        logger.info(f"添加扩展名找到文件: {candidate}")
                        return candidate
        
        # 4. 部分匹配：文件名包含在文件名的开头（忽略扩展名）
        # 先进行部分匹配，因为用户可能只输入文件名的一部分（如"强制执行申请书"匹配"强制执行申请书.pdf"）
        filename_without_ext = Path(filename).stem  # 获取不带扩展名的文件名
        filename_without_ext_lower = filename_without_ext.lower()
        
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            try:
                for item in search_dir.iterdir():
                    if item.is_file():
                        item_stem = item.stem  # 获取文件名的stem部分（不含扩展名）
                        item_stem_lower = item_stem.lower()
                        # 如果用户输入的文件名完全匹配文件名的stem（大小写不敏感），或者文件名的stem包含用户输入
                        if (filename_without_ext_lower == item_stem_lower or 
                            filename_without_ext_lower in item_stem_lower or 
                            item_stem_lower.startswith(filename_without_ext_lower) or
                            item_stem_lower == filename_without_ext_lower):
                            logger.info(f"部分匹配找到文件: {item} (匹配: {filename_without_ext})")
                            return item
            except (PermissionError, OSError):
                continue
        
        # 5. 递归搜索（限制深度为2，仅在前面的搜索失败时使用）
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            try:
                # 先尝试精确匹配（大小写敏感）
                for item in search_dir.rglob(filename):
                    if item.is_file():
                        logger.info(f"递归搜索找到文件: {item}")
                        return item
                
                # 再尝试精确匹配（大小写不敏感）
                for item in search_dir.rglob("*"):
                    if item.is_file() and item.name.lower() == filename_lower:
                        logger.info(f"递归搜索（大小写不敏感）找到文件: {item}")
                        return item
                
                # 再尝试部分匹配（递归）
                filename_without_ext = Path(filename).stem
                filename_without_ext_lower = filename_without_ext.lower()
                for item in search_dir.rglob("*"):
                    if item.is_file():
                        item_stem = item.stem
                        item_stem_lower = item_stem.lower()
                        # 完全匹配或包含匹配（大小写不敏感）
                        if (filename_without_ext_lower == item_stem_lower or 
                            filename_without_ext_lower in item_stem_lower or 
                            item_stem_lower.startswith(filename_without_ext_lower)):
                            logger.info(f"递归部分匹配找到文件: {item} (匹配: {filename_without_ext})")
                            return item
            except (PermissionError, OSError):
                continue
        
        logger.warning(f"未找到文件: {filename}（已搜索: {', '.join([str(d) for d in search_dirs if d.exists()])}）")
        return None
    
    def _find_folder(self, folder_name: str, search_dirs: List[Path] = None) -> Optional[Path]:
        """
        智能搜索文件夹：在指定目录中搜索文件夹名
        
        支持：
        - 精确匹配（如 "8888"）
        - 部分匹配（如 "88" 可以匹配 "8888"）
        
        Args:
            folder_name: 文件夹名（如 "8888"）
            search_dirs: 搜索目录列表，默认为用户主目录下的常用目录
        
        Returns:
            找到的文件夹路径，如果未找到返回None
        """
        if search_dirs is None:
            home = Path.home()
            search_dirs = [
                home / "Desktop",
                home / "Downloads",
                home / "Documents",
                home,  # 也在主目录下搜索
                self.sandbox_path,
            ]
        
        # 规范化文件夹名（去除前后空格）
        folder_name = folder_name.strip()
        
        # 1. 先尝试直接匹配文件夹名（精确匹配）
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            candidate = search_dir / folder_name
            if candidate.exists() and candidate.is_dir():
                logger.info(f"精确匹配找到文件夹: {candidate}")
                return candidate
        
        # 2. 部分匹配：文件夹名包含在文件夹名的开头
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            try:
                for item in search_dir.iterdir():
                    if item.is_dir():
                        item_name = item.name
                        # 如果用户输入的文件夹名包含在文件夹名的开头，或者文件夹名包含用户输入
                        if folder_name in item_name or item_name.startswith(folder_name):
                            logger.info(f"部分匹配找到文件夹: {item} (匹配: {folder_name})")
                            return item
            except (PermissionError, OSError):
                continue
        
        # 3. 递归搜索（限制深度为2）
        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            try:
                # 先尝试精确匹配
                for item in search_dir.rglob(folder_name):
                    if item.is_dir():
                        logger.info(f"递归搜索找到文件夹: {item}")
                        return item
                
                # 再尝试部分匹配（递归）
                for item in search_dir.rglob("*"):
                    if item.is_dir():
                        item_name = item.name
                        if folder_name in item_name or item_name.startswith(folder_name):
                            logger.info(f"递归部分匹配找到文件夹: {item} (匹配: {folder_name})")
                            return item
            except (PermissionError, OSError):
                continue
        
        logger.warning(f"未找到文件夹: {folder_name}")
        return None
    
    def _rename(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        重命名文件（支持智能文件搜索）
        
        如果file_path只是文件名（不包含路径），会自动在用户主目录下搜索
        """
        file_path_str = params.get("file_path", "")
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        
        new_name = params.get("new_name", "")
        if not new_name:
            raise FileManagerError("重命名参数缺少new_name")
        
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名（不包含路径分隔符），尝试智能搜索
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            logger.info(f"检测到文件名格式，开始智能搜索: {file_path_str}")
            found_file = self._find_file(file_path_str)
            if found_file:
                file_path = found_file
                logger.info(f"找到文件: {file_path}")
            else:
                raise FileManagerError(
                    f"未找到文件: {file_path_str}。"
                    f"请提供完整路径，如 '~/Desktop/{file_path_str}' 或 '/Users/username/Desktop/{file_path_str}'"
                )
        # 如果是相对路径（包含路径分隔符但不是绝对路径），相对于用户主目录
        elif not file_path.is_absolute() and ("/" in file_path_str or "\\" in file_path_str):
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        if not file_path.exists():
            raise FileManagerError(f"文件不存在: {file_path}")
        
        if not file_path.is_file():
            raise FileManagerError(f"路径不是文件: {file_path}")
        
        # 安全：验证新名称不包含路径分隔符
        if "/" in new_name or "\\" in new_name:
            raise FileManagerError(f"新名称不能包含路径分隔符: {new_name}")
        
        target_path = file_path.parent / new_name
        
        # 如果目标文件已存在，添加序号
        if target_path.exists():
            counter = 1
            stem = target_path.stem
            suffix = target_path.suffix
            while target_path.exists():
                target_path = file_path.parent / f"{stem}_{counter}{suffix}"
                counter += 1
        
        # 安全：使用安全的移动方法（在同一目录内重命名）
        self._safe_move(file_path, target_path)
        
        return {
            "success": True,
            "message": f"文件已重命名: {target_path}",
            "data": {"source": str(file_path), "target": str(target_path)}
        }
    
    def _move(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        移动文件（支持智能文件搜索和 ~ 符号）
        
        如果file_path只是文件名（不包含路径），会自动在用户主目录下搜索
        """
        file_path_str = params.get("file_path", "")
        target_dir_str = params.get("target_dir", "")
        
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        if not target_dir_str:
            raise FileManagerError("缺少target_dir参数")
        
        # ========== 处理源文件路径 ==========
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名（不包含路径分隔符），尝试智能搜索
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            logger.info(f"检测到文件名格式，开始智能搜索: {file_path_str}")
            found_file = self._find_file(file_path_str)
            if found_file:
                file_path = found_file
                logger.info(f"找到文件: {file_path}")
            else:
                raise FileManagerError(
                    f"未找到文件: {file_path_str}。"
                    f"请提供完整路径，如 '~/Desktop/{file_path_str}' 或 '/Users/username/Desktop/{file_path_str}'"
                )
        # 如果是相对路径（包含路径分隔符但不是绝对路径），相对于用户主目录
        elif not file_path.is_absolute() and ("/" in file_path_str or "\\" in file_path_str):
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 安全：验证源文件路径
        file_path = self._validate_path(file_path)
        
        if not file_path.exists():
            raise FileManagerError(f"文件不存在: {file_path}")
        
        if not file_path.is_file():
            raise FileManagerError(f"路径不是文件: {file_path}")
        
        # ========== 处理目标目录路径 ==========
        # 处理 ~ 符号
        if target_dir_str.startswith("~/"):
            target_dir_str = str(Path.home() / target_dir_str[2:])
        elif target_dir_str.startswith("~"):
            target_dir_str = str(Path.home() / target_dir_str[1:])
        
        target_dir = Path(target_dir_str)
        
        # 如果只是文件夹名（不包含路径分隔符），尝试智能搜索
        if "/" not in target_dir_str and "\\" not in target_dir_str and not target_dir_str.startswith("~"):
            logger.info(f"检测到目标文件夹名格式，开始智能搜索: {target_dir_str}")
            found_folder = self._find_folder(target_dir_str)
            if found_folder:
                target_dir = found_folder
                logger.info(f"找到目标文件夹: {target_dir}")
            else:
                raise FileManagerError(
                    f"未找到目标文件夹: {target_dir_str}。"
                    f"请提供完整路径，如 '~/Desktop/{target_dir_str}' 或 '/Users/username/Desktop/{target_dir_str}'"
                )
        # 如果是相对路径（包含路径分隔符但不是绝对路径），相对于用户主目录
        elif not target_dir.is_absolute() and ("/" in target_dir_str or "\\" in target_dir_str):
            target_dir = Path.home() / target_dir
        
        target_dir = target_dir.resolve()
        
        # 安全：验证目标目录（使用相同的验证逻辑）
        try:
            self._validate_path(target_dir)
        except FileManagerError as e:
            raise FileManagerError(
                f"目标目录不在允许的操作范围内: {target_dir}。"
                f"只允许操作用户主目录和沙盒目录下的文件。"
            )
        
        target_dir.mkdir(parents=True, exist_ok=True)
        target_path = target_dir / file_path.name
        
        # 如果目标文件已存在，添加序号
        if target_path.exists():
            counter = 1
            stem = file_path.stem
            suffix = file_path.suffix
            while target_path.exists():
                target_path = target_dir / f"{stem}_{counter}{suffix}"
                counter += 1
        
        # 安全：使用安全的移动方法
        self._safe_move(file_path, target_path)
        
        return {
            "success": True,
            "message": f"文件已移动到: {target_path}",
            "data": {"source": str(file_path), "target": str(target_path)}
        }
    
    def _classify(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """分类文件（调用organize）"""
        return self._organize(params)
    
    def _batch_rename(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        批量重命名文件夹内的文件
        
        Args:
            params: 包含：
                - folder_path: 文件夹路径（必需），支持智能搜索（如 "8888"）
                - name_pattern: 命名模式（必需），如 "1234"（会生成 1234_1.png, 1234_2.jpg 等）
                - file_filter: 文件过滤条件（可选），如 "image"（只处理图片文件）
                - sort_by: 排序方式（可选），如 "size"（按文件大小排序）
        
        Returns:
            批量重命名结果
        """
        folder_path_str = params.get("folder_path", "")
        name_pattern = params.get("name_pattern", "")
        file_filter = params.get("file_filter", "")  # 如 "image", "pdf" 等
        sort_by = params.get("sort_by", "")  # 如 "size", "name", "date"
        use_simple_numbering = params.get("use_simple_numbering", False)
        
        if not folder_path_str:
            raise FileManagerError("缺少folder_path参数")
        
        # 如果 name_pattern 为空，默认使用纯数字命名
        if not name_pattern:
            use_simple_numbering = True
        
        # ========== 处理文件夹路径 ==========
        # 处理 ~ 符号
        if folder_path_str.startswith("~/"):
            folder_path_str = str(Path.home() / folder_path_str[2:])
        elif folder_path_str.startswith("~"):
            folder_path_str = str(Path.home() / folder_path_str[1:])
        
        folder_path = Path(folder_path_str)
        
        # 如果只是文件夹名（不包含路径分隔符），尝试智能搜索
        if "/" not in folder_path_str and "\\" not in folder_path_str and not folder_path_str.startswith("~"):
            logger.info(f"检测到文件夹名格式，开始智能搜索: {folder_path_str}")
            found_folder = self._find_folder(folder_path_str)
            if found_folder:
                folder_path = found_folder
                logger.info(f"找到文件夹: {folder_path}")
            else:
                raise FileManagerError(
                    f"未找到文件夹: {folder_path_str}。"
                    f"请提供完整路径，如 '~/Desktop/{folder_path_str}' 或 '/Users/username/Desktop/{folder_path_str}'"
                )
        # 如果是相对路径，相对于用户主目录
        elif not folder_path.is_absolute():
            folder_path = Path.home() / folder_path
        
        folder_path = folder_path.resolve()
        
        # 安全：验证路径
        folder_path = self._validate_path(folder_path)
        
        if not folder_path.exists():
            raise FileManagerError(f"文件夹不存在: {folder_path}")
        
        if not folder_path.is_dir():
            raise FileManagerError(f"路径不是文件夹: {folder_path}")
        
        # ========== 获取文件列表 ==========
        files = []
        image_extensions = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg", ".ico"}
        pdf_extensions = {".pdf"}
        
        for item in folder_path.iterdir():
            if item.is_file():
                # 文件过滤
                if file_filter:
                    if file_filter.lower() == "image":
                        if item.suffix.lower() not in image_extensions:
                            continue
                    elif file_filter.lower() == "pdf":
                        if item.suffix.lower() not in pdf_extensions:
                            continue
                    # 可以扩展更多过滤条件
                
                files.append(item)
        
        if not files:
            raise FileManagerError(f"文件夹中没有符合条件的文件: {folder_path}")
        
        # ========== 排序 ==========
        if sort_by == "size":
            files.sort(key=lambda f: f.stat().st_size)
        elif sort_by == "name":
            files.sort(key=lambda f: f.name.lower())
        elif sort_by == "date":
            files.sort(key=lambda f: f.stat().st_mtime)
        # 默认按文件名排序
        
        # ========== 批量重命名 ==========
        renamed_count = 0
        renamed_files = []
        
        # 判断是否使用纯数字命名
        # 如果 name_pattern 为空或只有数字（且长度<=2），且用户没有明确要求前缀格式，使用纯数字
        if not name_pattern or (name_pattern.isdigit() and len(name_pattern) <= 2):
            use_simple_numbering = True
        
        for index, file_path in enumerate(files, start=1):
            # 生成新文件名
            if use_simple_numbering:
                # 纯数字模式：1.png, 2.jpg, 3.png 等
                new_name = f"{index}{file_path.suffix}"
            else:
                # 前缀模式：name_pattern_index.extension
                new_name = f"{name_pattern}_{index}{file_path.suffix}"
            target_path = folder_path / new_name
            
            # 如果目标文件已存在，跳过（避免覆盖）
            if target_path.exists() and target_path != file_path:
                logger.warning(f"目标文件已存在，跳过: {target_path}")
                continue
            
            # 重命名
            try:
                file_path.rename(target_path)
                renamed_count += 1
                renamed_files.append({
                    "old_name": file_path.name,
                    "new_name": new_name
                })
                logger.info(f"✅ 重命名: {file_path.name} -> {new_name}")
            except Exception as e:
                logger.error(f"重命名失败: {file_path.name} -> {new_name}, 错误: {e}")
                continue
        
        return {
            "success": True,
            "message": f"批量重命名完成: {renamed_count}/{len(files)} 个文件已重命名",
            "data": {
                "folder": str(folder_path),
                "renamed_count": renamed_count,
                "total_count": len(files),
                "renamed_files": renamed_files
            }
        }
    
    def _copy(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        复制单个文件（创建副本）
        
        Args:
            params: 包含：
                - file_path: 源文件路径（必需），支持：
                  * 相对路径：如 "Desktop/file.png"（相对于用户主目录）
                  * 绝对路径：如 "~/Desktop/file.png" 或 "/Users/username/Desktop/file.png"
                  * 文件名：如 "screenshot.png"（会在用户主目录下搜索）
                  * **如果上下文中有文件路径，优先使用上下文中的路径**（更准确）
                - target_path: 目标文件路径（可选），支持：
                  * 相对路径：如 "Desktop/file_copy.png"
                  * 绝对路径：如 "~/Desktop/file_copy.png"
                  * 如果只提供目录（如 "~/Desktop"），会自动使用源文件名并添加 "_copy" 后缀
                  * 如果不提供，默认在源文件同目录下创建副本，并添加 "_copy" 后缀
                - copy_suffix: 副本文件名后缀（可选），如 "_backup"（会生成 file_backup.png），默认为 "_copy"
        
        Returns:
            复制结果
        """
        file_path_str = params.get("file_path")
        target_path_str = params.get("target_path")
        copy_suffix = params.get("copy_suffix", "_copy")
        
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        
        # ========== 处理源文件路径 ==========
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名（不包含路径分隔符），尝试智能搜索
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            logger.info(f"检测到文件名格式，开始智能搜索: {file_path_str}")
            found_file = self._find_file(file_path_str)
            if found_file:
                file_path = found_file
                logger.info(f"找到文件: {file_path}")
            else:
                raise FileManagerError(
                    f"未找到文件: {file_path_str}。"
                    f"请提供完整路径，如 '~/Desktop/{file_path_str}' 或 '/Users/username/Desktop/{file_path_str}'"
                )
        # 如果是相对路径，相对于用户主目录
        elif not file_path.is_absolute():
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        # 检查源文件是否存在
        if not file_path.exists():
            raise FileManagerError(f"文件不存在: {file_path}")
        
        if not file_path.is_file():
            raise FileManagerError(f"路径不是文件: {file_path}")
        
        # ========== 处理目标路径 ==========
        # 兼容 destination 参数（Planner 可能使用这个名称）
        if not target_path_str:
            target_path_str = params.get("destination")
        
        if target_path_str:
            # 处理 ~ 符号
            if target_path_str.startswith("~/"):
                target_path_str = str(Path.home() / target_path_str[2:])
            elif target_path_str.startswith("~"):
                target_path_str = str(Path.home() / target_path_str[1:])
            
            target_path = Path(target_path_str)
            
            # 智能搜索：如果只是文件夹名（不包含路径分隔符），尝试搜索文件夹
            if "/" not in target_path_str and "\\" not in target_path_str and not target_path_str.startswith("~"):
                logger.info(f"检测到目标文件夹名格式，开始智能搜索: {target_path_str}")
                found_folder = self._find_folder(target_path_str)
                if found_folder:
                    target_path = found_folder
                    logger.info(f"找到目标文件夹: {target_path}")
                    # 目标路径是文件夹，使用源文件名并添加后缀
                    stem = file_path.stem
                    suffix = file_path.suffix
                    target_path = target_path / f"{stem}{copy_suffix}{suffix}"
                else:
                    raise FileManagerError(
                        f"未找到目标文件夹: {target_path_str}。"
                        f"请提供完整路径，如 '~/Desktop/{target_path_str}' 或 '/Users/username/Desktop/{target_path_str}'"
                    )
            # 如果是相对路径，相对于用户主目录
            elif not target_path.is_absolute():
                target_path = Path.home() / target_path
            
            target_path = target_path.resolve()
            
            # 安全：验证目标路径
            target_path = self._validate_path(target_path)
            
            # 如果目标路径是目录，使用源文件名并添加后缀
            if target_path.is_dir() or (not target_path.exists() and target_path_str.endswith("/")):
                stem = file_path.stem
                suffix = file_path.suffix
                target_path = target_path / f"{stem}{copy_suffix}{suffix}"
            # 如果目标路径没有扩展名，使用源文件的扩展名
            elif not target_path.suffix and file_path.suffix:
                target_path = target_path.with_suffix(file_path.suffix)
        else:
            # 如果没有提供目标路径，默认在源文件同目录下创建副本
            stem = file_path.stem
            suffix = file_path.suffix
            target_path = file_path.parent / f"{stem}{copy_suffix}{suffix}"
        
        # 如果目标文件已存在，添加序号
        if target_path.exists():
            counter = 1
            original_target = target_path
            while target_path.exists():
                stem = original_target.stem
                suffix = original_target.suffix
                # 移除可能已存在的序号后缀
                if copy_suffix in stem:
                    base_stem = stem.rsplit(copy_suffix, 1)[0]
                else:
                    base_stem = stem
                target_path = original_target.parent / f"{base_stem}{copy_suffix}_{counter}{suffix}"
                counter += 1
        
        # 确保目标目录存在
        target_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 复制文件
        try:
            shutil.copy2(file_path, target_path)
            logger.info(f"✅ 已复制文件: {file_path} -> {target_path}")
            
            return {
                "success": True,
                "message": f"已复制文件: {target_path}",
                "data": {
                    "source": str(file_path),
                    "target": str(target_path),
                    "new_path": str(target_path)  # 用于上下文更新
                }
            }
        except Exception as e:
            error_msg = f"复制文件失败: {e}"
            logger.error(error_msg, exc_info=True)
            raise FileManagerError(error_msg)
    
    def _batch_copy(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        批量复制文件夹内的文件（创建副本）
        
        Args:
            params: 包含：
                - folder_path: 文件夹路径（必需），支持智能搜索（如 "8888"）
                - file_filter: 文件过滤条件（可选），如 "image"（只处理图片文件）
                - copy_suffix: 副本文件名后缀（可选），如 "_copy"（会生成 file_copy.png），默认为 "_copy"
        
        Returns:
            批量复制结果
        """
        folder_path_str = params.get("folder_path", "")
        file_filter = params.get("file_filter", "")  # 如 "image", "pdf" 等
        copy_suffix = params.get("copy_suffix", "_copy")  # 副本后缀，默认为 "_copy"
        
        if not folder_path_str:
            raise FileManagerError("缺少folder_path参数")
        
        # ========== 处理文件夹路径 ==========
        # 处理 ~ 符号
        if folder_path_str.startswith("~/"):
            folder_path_str = str(Path.home() / folder_path_str[2:])
        elif folder_path_str.startswith("~"):
            folder_path_str = str(Path.home() / folder_path_str[1:])
        
        folder_path = Path(folder_path_str)
        
        # 如果只是文件夹名（不包含路径分隔符），尝试智能搜索
        if "/" not in folder_path_str and "\\" not in folder_path_str and not folder_path_str.startswith("~"):
            logger.info(f"检测到文件夹名格式，开始智能搜索: {folder_path_str}")
            found_folder = self._find_folder(folder_path_str)
            if found_folder:
                folder_path = found_folder
                logger.info(f"找到文件夹: {folder_path}")
            else:
                raise FileManagerError(
                    f"未找到文件夹: {folder_path_str}。"
                    f"请提供完整路径，如 '~/Desktop/{folder_path_str}' 或 '/Users/username/Desktop/{folder_path_str}'"
                )
        # 如果是相对路径，相对于用户主目录
        elif not folder_path.is_absolute():
            folder_path = Path.home() / folder_path
        
        folder_path = folder_path.resolve()
        
        # 安全：验证路径
        folder_path = self._validate_path(folder_path)
        
        if not folder_path.exists():
            raise FileManagerError(f"文件夹不存在: {folder_path}")
        
        if not folder_path.is_dir():
            raise FileManagerError(f"路径不是文件夹: {folder_path}")
        
        # ========== 获取文件列表 ==========
        files = []
        image_extensions = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg", ".ico"}
        pdf_extensions = {".pdf"}
        
        for item in folder_path.iterdir():
            if item.is_file():
                # 文件过滤
                if file_filter:
                    if file_filter.lower() == "image":
                        if item.suffix.lower() not in image_extensions:
                            continue
                    elif file_filter.lower() == "pdf":
                        if item.suffix.lower() not in pdf_extensions:
                            continue
                    # 可以扩展更多过滤条件
                
                files.append(item)
        
        if not files:
            raise FileManagerError(f"文件夹中没有符合条件的文件: {folder_path}")
        
        # ========== 批量复制 ==========
        copied_count = 0
        copied_files = []
        
        for file_path in files:
            # 生成副本文件名：原文件名_copy.extension
            stem = file_path.stem
            suffix = file_path.suffix
            new_name = f"{stem}{copy_suffix}{suffix}"
            target_path = folder_path / new_name
            
            # 如果目标文件已存在，添加序号
            if target_path.exists():
                counter = 1
                while target_path.exists():
                    new_name = f"{stem}{copy_suffix}_{counter}{suffix}"
                    target_path = folder_path / new_name
                    counter += 1
            
            # 复制文件
            try:
                shutil.copy2(file_path, target_path)
                copied_count += 1
                copied_files.append({
                    "original": file_path.name,
                    "copy": new_name
                })
                logger.info(f"✅ 已创建副本: {file_path.name} -> {new_name}")
            except Exception as e:
                logger.error(f"复制失败: {file_path.name} -> {new_name}, 错误: {e}")
                continue
        
        return {
            "success": True,
            "message": f"批量复制完成: {copied_count}/{len(files)} 个文件已创建副本",
            "data": {
                "folder": str(folder_path),
                "copied_count": copied_count,
                "total_count": len(files),
                "copied_files": copied_files
            }
        }
    
    def _create_file(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        创建新文件
        
        Args:
            params: 包含：
                - file_path: 文件路径（必需），支持：
                  * 相对路径：如 "Desktop/document.docx"（相对于用户主目录）
                  * 绝对路径：如 "~/Desktop/document.docx" 或 "/Users/username/Desktop/document.docx"
                - content: 文件内容（可选），如果不提供，创建空文件
                - file_type: 文件类型（可选），如 "word"、"text"、"excel" 等，用于生成合适的文件扩展名
        
        Returns:
            创建结果
        """
        file_path_str = params.get("file_path", "")
        content = params.get("content", "")
        file_type = params.get("file_type", "").lower()
        
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        
        # ========== 处理文件路径 ==========
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名（不包含路径分隔符），默认保存到桌面
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            logger.info(f"检测到文件名格式，默认保存到桌面: {file_path_str}")
            file_path = Path.home() / "Desktop" / file_path_str
        
        # 如果是相对路径，相对于用户主目录
        elif not file_path.is_absolute():
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 如果文件路径没有扩展名，根据 file_type 添加
        if not file_path.suffix:
            type_extensions = {
                "word": ".docx",
                "doc": ".docx",
                "docx": ".docx",
                "text": ".txt",
                "txt": ".txt",
                "excel": ".xlsx",
                "xls": ".xlsx",
                "xlsx": ".xlsx",
                "pdf": ".pdf",
                "markdown": ".md",
                "md": ".md",
            }
            if file_type and file_type in type_extensions:
                file_path = file_path.with_suffix(type_extensions[file_type])
            else:
                # 默认使用 .txt
                file_path = file_path.with_suffix(".txt")
        
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        # 确保父目录存在
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 如果文件已存在，添加序号
        original_path = file_path
        if file_path.exists():
            counter = 1
            stem = file_path.stem
            suffix = file_path.suffix
            while file_path.exists():
                file_path = original_path.parent / f"{stem}_{counter}{suffix}"
                counter += 1
            logger.info(f"文件已存在，使用新名称: {file_path}")
        
        # 创建文件
        try:
            if content:
                # 写入内容
                file_path.write_text(content, encoding="utf-8")
            else:
                # 创建空文件
                file_path.touch()
            
            logger.info(f"✅ 已创建文件: {file_path}")
            
            return {
                "success": True,
                "message": f"已创建文件: {file_path}",
                "data": {
                    "path": str(file_path),
                    "size": file_path.stat().st_size if file_path.exists() else 0
                }
            }
            
        except Exception as e:
            logger.error(f"创建文件失败: {e}", exc_info=True)
            raise FileManagerError(f"创建文件失败: {e}")
    
    def _read_file(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        读取文件内容
        
        Args:
            params: 包含：
                - file_path: 文件路径（必需），支持：
                  * 相对路径：如 "Desktop/file.txt"（相对于用户主目录）
                  * 绝对路径：如 "~/Desktop/file.txt" 或 "/Users/username/Desktop/file.txt"
                  * 文件名：如 "test.txt"（会在用户主目录下搜索）
        
        Returns:
            读取结果，包含文件内容
        """
        file_path_str = params.get("file_path", "")
        
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        
        # ========== 处理文件路径 ==========
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名（不包含路径分隔符），尝试智能搜索
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            logger.info(f"检测到文件名格式，开始智能搜索: {file_path_str}")
            found_file = self._find_file(file_path_str)
            if found_file:
                file_path = found_file
                logger.info(f"找到文件: {file_path}")
            else:
                raise FileManagerError(f"文件不存在: {file_path_str}")
        
        # 如果是相对路径，相对于用户主目录
        elif not file_path.is_absolute():
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        if not file_path.exists():
            raise FileManagerError(f"文件不存在: {file_path}")
        
        if not file_path.is_file():
            raise FileManagerError(f"路径不是文件: {file_path}")
        
        # 读取文件
        try:
            # 检查是否是 Word 文档
            file_ext = file_path.suffix.lower()
            if file_ext in [".docx"]:
                # 尝试使用 python-docx 读取 Word 文档
                try:
                    from docx import Document
                    doc = Document(str(file_path))
                    # 提取所有段落文本（保留完整内容）
                    paragraphs = []
                    for para in doc.paragraphs:
                        paragraphs.append(para.text)
                    content = "\n".join(paragraphs)
                    
                    # 保存文档对象和段落对象，以便后续修改时保留格式
                    logger.info(f"✅ 已读取 Word 文档: {file_path} ({len(content)} 字符, {len(paragraphs)} 段落)")
                    
                    return {
                        "success": True,
                        "message": f"已读取 Word 文档: {file_path}",
                        "data": {
                            "path": str(file_path),
                            "content": content,
                            "size": len(content),
                            "file_type": "docx",
                            "paragraphs": paragraphs,  # 保留段落信息，方便后续修改
                            "original_file_path": str(file_path)  # 保存原始文件路径，用于后续修改
                        }
                    }
                except ImportError:
                    logger.warning("python-docx 库未安装，无法读取 Word 文档，尝试作为文本文件读取")
                    # 如果没有安装 python-docx，继续尝试文本读取
                except Exception as e:
                    logger.warning(f"使用 python-docx 读取失败: {e}，尝试作为文本文件读取")
            
            # 普通文本文件读取
            content = file_path.read_text(encoding="utf-8")
            logger.info(f"✅ 已读取文件: {file_path} ({len(content)} 字符)")
            
            return {
                "success": True,
                "message": f"已读取文件: {file_path}",
                "data": {
                    "path": str(file_path),
                    "content": content,
                    "size": len(content)
                }
            }
            
        except UnicodeDecodeError:
            # 尝试二进制读取
            try:
                content_bytes = file_path.read_bytes()
                logger.warning(f"文件包含非UTF-8字符，返回二进制数据: {file_path}")
                return {
                    "success": True,
                    "message": f"已读取文件（二进制）: {file_path}",
                    "data": {
                        "path": str(file_path),
                        "content": content_bytes.hex(),  # 返回十六进制字符串
                        "size": len(content_bytes),
                        "is_binary": True
                    }
                }
            except Exception as e:
                logger.error(f"读取文件失败: {e}", exc_info=True)
                raise FileManagerError(f"读取文件失败: {e}")
        except Exception as e:
            logger.error(f"读取文件失败: {e}", exc_info=True)
            raise FileManagerError(f"读取文件失败: {e}")
    
    def _write_file(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        写入内容到文件（支持追加或覆盖）
        
        Args:
            params: 包含：
                - file_path: 文件路径（必需），支持：
                  * 相对路径：如 "Desktop/file.txt"（相对于用户主目录）
                  * 绝对路径：如 "~/Desktop/file.txt" 或 "/Users/username/Desktop/file.txt"
                  * 文件名：如 "test.txt"（会在用户主目录下搜索）
                - content: 文件内容（必需）
                - append: 是否追加模式（可选，布尔值），如果为 true，追加到文件末尾；如果为 false 或未指定，覆盖文件
                - newline: 是否在追加时添加换行符（可选，布尔值），默认为 true
        
        Returns:
            写入结果
        """
        file_path_str = params.get("file_path", "")
        content = params.get("content", "")
        append = params.get("append", False)
        newline = params.get("newline", True)
        
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        
        if content is None or content == "":
            raise FileManagerError("缺少content参数")
        
        # ========== 处理文件路径 ==========
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名（不包含路径分隔符），尝试智能搜索
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            logger.info(f"检测到文件名格式，开始智能搜索: {file_path_str}")
            found_file = self._find_file(file_path_str)
            if found_file:
                file_path = found_file
                logger.info(f"找到文件: {file_path}")
            else:
                # 如果没找到文件，抛出错误（不要自动创建新文件）
                raise FileManagerError(
                    f"未找到文件: {file_path_str}。"
                    f"请提供完整路径，如 '~/Desktop/{file_path_str}' 或 '/Users/username/Desktop/{file_path_str}'"
                )
        
        # 如果是相对路径，相对于用户主目录
        elif not file_path.is_absolute():
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        # 确保父目录存在
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 如果文件不存在且不是追加模式，创建新文件
        if not file_path.exists() and not append:
            logger.info(f"文件不存在，将创建新文件: {file_path}")
        
        # 写入文件
        try:
            # 检查是否是 Word 文档
            file_ext = file_path.suffix.lower()
            if file_ext in [".docx"]:
                # 尝试使用 python-docx 写入 Word 文档
                try:
                    from docx import Document
                    
                    content_str = str(content)
                    
                    # 如果文件存在，打开现有文档；否则创建新文档
                    if file_path.exists():
                        doc = Document(str(file_path))
                        
                        if append:
                            # 追加模式：在现有文档末尾添加新段落
                            paragraphs = content_str.split("\n")
                            for para_text in paragraphs:
                                if para_text.strip():
                                    doc.add_paragraph(para_text)
                        else:
                            # 覆盖模式：智能替换
                            # 如果新内容看起来像是完整的文档（多段落），则替换所有段落
                            # 如果新内容看起来像是部分内容（单段落或短文本），则尝试在现有段落中查找并替换
                            new_paragraphs = [p.strip() for p in content_str.split("\n") if p.strip()]
                            existing_paragraphs = list(doc.paragraphs)
                            
                            # 检查新内容是否包含完整的文档结构（多段落）
                            if len(new_paragraphs) >= len(existing_paragraphs) * 0.8:
                                # 新内容看起来像是完整文档，替换所有段落
                                # 先清空现有段落内容
                                for para in existing_paragraphs:
                                    para.text = ""
                                # 然后添加新内容
                                for i, para_text in enumerate(new_paragraphs):
                                    if i < len(existing_paragraphs):
                                        existing_paragraphs[i].text = para_text
                                    else:
                                        doc.add_paragraph(para_text)
                                # 删除多余的段落
                                if len(new_paragraphs) < len(existing_paragraphs):
                                    for i in range(len(existing_paragraphs) - 1, len(new_paragraphs) - 1, -1):
                                        p = existing_paragraphs[i]._element
                                        p.getparent().remove(p)
                            else:
                                # 新内容看起来像是部分内容，尝试在现有段落中查找并替换
                                # 将新内容合并为单个字符串，然后在每个段落中查找并替换
                                new_content_combined = " ".join(new_paragraphs)
                                
                                # 尝试在每个段落中查找并替换
                                replaced = False
                                for para in existing_paragraphs:
                                    if new_content_combined in para.text or para.text in new_content_combined:
                                        # 如果新内容包含在段落中，或者段落包含在新内容中，进行替换
                                        para.text = new_content_combined
                                        replaced = True
                                        break
                                
                                if not replaced:
                                    # 如果没有找到匹配的段落，尝试在整个文档中查找并替换
                                    # 将整个文档内容合并，查找并替换
                                    full_text = "\n".join([p.text for p in existing_paragraphs])
                                    if "XXXXXXXXXXX" in full_text or "XXXX" in full_text:
                                        # 在第一个包含目标文本的段落中替换
                                        for para in existing_paragraphs:
                                            if "XXXXXXXXXXX" in para.text or "XXXX" in para.text:
                                                para.text = para.text.replace("XXXXXXXXXXX", new_content_combined).replace("XXXX", new_content_combined)
                                                replaced = True
                                                break
                                
                                if not replaced:
                                    # 如果还是没有找到匹配，追加到文档末尾
                                    doc.add_paragraph(new_content_combined)
                    else:
                        # 文件不存在，创建新文档
                        doc = Document()
                        paragraphs = content_str.split("\n")
                        for para_text in paragraphs:
                            if para_text.strip():
                                doc.add_paragraph(para_text)
                    
                    doc.save(str(file_path))
                    logger.info(f"✅ 已{'追加' if append else '写入'}内容到 Word 文档: {file_path}")
                    
                    return {
                        "success": True,
                        "message": f"已{'追加' if append else '写入'}内容到 Word 文档: {file_path}",
                        "data": {
                            "path": str(file_path),
                            "size": file_path.stat().st_size if file_path.exists() else 0,
                            "append": append,
                            "file_type": "docx"
                        }
                    }
                except ImportError:
                    logger.warning("python-docx 库未安装，无法写入 Word 文档，尝试作为文本文件写入")
                    # 如果没有安装 python-docx，继续尝试文本写入
                except Exception as e:
                    logger.warning(f"使用 python-docx 写入失败: {e}，尝试作为文本文件写入")
            
            # 普通文本文件写入
            if append:
                # 追加模式
                with open(file_path, "a", encoding="utf-8") as f:
                    if newline and file_path.exists() and file_path.stat().st_size > 0:
                        f.write("\n")
                    f.write(str(content))
                logger.info(f"✅ 已追加内容到文件: {file_path}")
            else:
                # 覆盖模式
                file_path.write_text(str(content), encoding="utf-8")
                logger.info(f"✅ 已写入内容到文件: {file_path}")
            
            return {
                "success": True,
                "message": f"已{'追加' if append else '写入'}内容到文件: {file_path}",
                "data": {
                    "path": str(file_path),
                    "size": file_path.stat().st_size if file_path.exists() else 0,
                    "append": append
                }
            }
            
        except Exception as e:
            logger.error(f"写入文件失败: {e}", exc_info=True)
            raise FileManagerError(f"写入文件失败: {e}")
    
    def _delete_file(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        删除文件或文件夹
        
        Args:
            params: 包含：
                - file_path: 文件/文件夹路径（必需）
        
        Returns:
            删除结果
        """
        file_path_str = params.get("file_path", "")
        
        if not file_path_str:
            raise FileManagerError("缺少file_path参数")
        
        # 处理 ~ 符号
        if file_path_str.startswith("~/"):
            file_path_str = str(Path.home() / file_path_str[2:])
        elif file_path_str.startswith("~"):
            file_path_str = str(Path.home() / file_path_str[1:])
        
        file_path = Path(file_path_str)
        
        # 如果只是文件名，尝试智能搜索
        if "/" not in file_path_str and "\\" not in file_path_str and not file_path_str.startswith("~"):
            found_file = self._find_file(file_path_str)
            if found_file:
                file_path = found_file
            else:
                raise FileManagerError(f"文件不存在: {file_path_str}")
        
        # 如果是相对路径，相对于用户主目录
        elif not file_path.is_absolute():
            file_path = Path.home() / file_path
        
        file_path = file_path.resolve()
        
        # 安全：验证路径
        file_path = self._validate_path(file_path)
        
        if not file_path.exists():
            raise FileManagerError(f"文件/文件夹不存在: {file_path}")
        
        try:
            if file_path.is_file():
                file_path.unlink()
                logger.info(f"✅ 已删除文件: {file_path}")
                return {
                    "success": True,
                    "message": "已删除文件: " + str(file_path),
                    "data": {"path": str(file_path), "type": "file"}
                }
            elif file_path.is_dir():
                shutil.rmtree(file_path)
                logger.info(f"✅ 已删除文件夹: {file_path}")
                return {
                    "success": True,
                    "message": "已删除文件夹: " + str(file_path),
                    "data": {"path": str(file_path), "type": "folder"}
                }
            else:
                raise FileManagerError(f"无法识别的文件类型: {file_path}")
        except Exception as e:
            logger.error(f"删除失败: {e}", exc_info=True)
            raise FileManagerError(f"删除失败: {e}")