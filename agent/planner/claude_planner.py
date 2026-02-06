"""
Claude AI规划器 - 简化重构版

核心理念：
- AI 自主思考，写代码解决问题
- 不再硬编码规则，让 AI 自己决定
- 只提供必要的能力描述和安全约束

遵循 docs/ARCHITECTURE.md 中的Planner模块规范
"""

import anthropic
import logging
import json
from typing import List, Dict, Any, Optional
from pathlib import Path

from agent.planner.base_planner import BasePlanner
from agent.tools.config import Config
from agent.tools.exceptions import PlannerError

logger = logging.getLogger(__name__)


class ClaudePlanner(BasePlanner):
    """
    Claude AI规划器 - 简化重构版
    
    核心改进：
    - 简化 prompt，不再硬编码工具规则
    - 让 AI 自主决定用工具还是写代码
    - 添加反思能力
    """
    
    def __init__(self, config: Config):
        """
        初始化Claude规划器
        
        Args:
            config: 配置对象
        """
        super().__init__(config)
        self.client = anthropic.Anthropic(api_key=config.api_key)
        self.model = config.model or "claude-sonnet-4-20250514"
        logger.info(f"Claude规划器已初始化，使用模型: {self.model}")
    
    def plan(
        self,
        user_instruction: str,
        context: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        规划任务：将用户指令转换为可执行步骤
        
        Args:
            user_instruction: 用户自然语言指令
            context: 上下文信息
        
        Returns:
            任务步骤列表
        """
        logger.info(f"开始规划任务: {user_instruction}")
        
        try:
            # 构建提示词
            prompt = self._build_prompt(user_instruction, context)
            
            def call_llm(user_prompt: str):
                return self.client.messages.create(
                    model=self.model,
                    max_tokens=4096,
                    messages=[{"role": "user", "content": user_prompt}],
                )

            # 调用Claude API
            logger.warning(f"🔵 正在调用Claude API规划任务...")
            response = call_llm(prompt)
            content = response.content[0].text if response.content else ""
            logger.warning(f"🔵 Claude原始响应（前2000字符）: {content[:2000]}...")
            logger.debug(f"Claude完整响应: {content}")
            
            # 解析响应：若 JSON 格式失败，自动重试一次（仅修复输出格式）
            try:
                steps = self._parse_response(content)
                logger.warning(f"🔵 解析后的步骤列表: {steps}")
                
                # 检查是否有open_app步骤，记录app_name用于调试
                for i, step in enumerate(steps):
                    if step.get("type") == "open_app":
                        app_name = step.get("params", {}).get("app_name", "")
                        logger.warning(f"🔵 步骤{i+1} open_app的app_name: '{app_name}' (长度: {len(app_name)})")
                        if len(app_name) > 20 or any(kw in app_name for kw in ["控制", "输入", "搜索", "按"]):
                            logger.error(f"❌ 检测到可疑的app_name: '{app_name}'，可能包含后续操作！AI没有正确拆分步骤！")
            except Exception as e:
                logger.warning(f"解析规划结果失败，将重试一次修复输出格式: {e}")
                retry_prompt = (
                    "你上一次的输出不是合法JSON，解析失败。\n"
                    "错误信息:\n"
                    + str(e)
                    + "\n\n"
                    "上一次原始输出（可能被截断）:\n"
                    + content[:1500]
                    + "\n\n"
                    "请重新输出合法JSON数组。规则：\n"
                    "- 只输出 JSON 数组（以 [ 开头，以 ] 结尾）\n"
                    "- 所有字符串必须使用双引号，且字符串内换行必须写成 \\n\n"
                    "- 不要输出 markdown 代码块\n"
                )
                response2 = call_llm(retry_prompt)
                content2 = response2.content[0].text if response2.content else ""
                logger.debug(f"Claude重试响应: {content2[:500]}...")
                steps = self._parse_response(content2)

            logger.info(f"规划完成，共 {len(steps)} 个步骤")
            
            # 后处理：修复截图工具选择错误
            user_instruction_lower = user_instruction.lower() if user_instruction else ""
            
            # 后处理1：检查是否有浏览器操作，如果有，后续的截图应该用 browser_screenshot 而不是 screenshot_desktop
            has_browser_operation = False
            browser_keywords = ["搜索", "打开网页", "访问", "浏览", "导航", "search", "navigate", "open", "visit", "browse"]
            for keyword in browser_keywords:
                if keyword in user_instruction:
                    has_browser_operation = True
                    break
            
            # 检查步骤中是否有浏览器操作
            for step in steps:
                step_type = step.get('type', '')
                if step_type in ['browser_navigate', 'browser_click', 'browser_input', 'browser_screenshot']:
                    has_browser_operation = True
                    break
            
            # 如果用户指令包含浏览器操作，检查是否有错误的 screenshot_desktop
            if has_browser_operation:
                for i, step in enumerate(steps, 1):
                    step_type = step.get('type')
                    step_params = step.get('params', {})
                    
                    # 如果发现 screenshot_desktop，但前面有浏览器操作，应该改为 browser_screenshot
                    if step_type == 'screenshot_desktop':
                        # 检查前面是否有浏览器操作步骤
                        has_browser_before = False
                        for j in range(i - 1):
                            prev_step_type = steps[j].get('type', '')
                            if prev_step_type in ['browser_navigate', 'browser_click', 'browser_input']:
                                has_browser_before = True
                                break
                        
                        # 如果前面有浏览器操作，或者用户指令明确包含浏览器操作，应该用 browser_screenshot
                        if has_browser_before or any(kw in user_instruction for kw in ["搜索", "打开", "访问", "浏览"]):
                            logger.warning(f"⚠️ 步骤 {i}：检测到浏览器操作后的截图，但使用了 screenshot_desktop，自动改为 browser_screenshot")
                            step['type'] = 'browser_screenshot'
                            # 如果 screenshot_desktop 有 save_path，保留它
                            if 'save_path' not in step_params or not step_params.get('save_path'):
                                # 检查用户是否要求保存到桌面
                                if "保存到桌面" in user_instruction or "保存桌面" in user_instruction or ("保存" in user_instruction and "桌面" in user_instruction):
                                    step_params['save_path'] = "~/Desktop/screenshot.png"
                                    step['params'] = step_params
                                    logger.info("✅ 已自动添加 save_path: ~/Desktop/screenshot.png")
                            logger.info(f"✅ 已自动将 screenshot_desktop 改为 browser_screenshot")
            
            # 后处理2：检查并修复 screenshot_desktop 缺少 save_path 的情况（仅当确实是桌面截图时）
            for i, step in enumerate(steps, 1):
                step_type = step.get('type')
                step_params = step.get('params', {})
                
                # 如果是 screenshot_desktop（且不是浏览器操作后的截图），检查用户是否要求保存到桌面
                if step_type == 'screenshot_desktop':
                    # 检查用户指令中是否包含"保存到桌面"、"保存桌面"等关键词
                    has_save_to_desktop = (
                        "保存到桌面" in user_instruction or
                        "保存桌面" in user_instruction or
                        "保存到 ~/Desktop" in user_instruction or
                        "save to desktop" in user_instruction_lower or
                        "save desktop" in user_instruction_lower or
                        ("保存" in user_instruction and "桌面" in user_instruction) or
                        ("save" in user_instruction_lower and "desktop" in user_instruction_lower)
                    )
                    
                    # 检查是否已经传递了 save_path 参数
                    has_save_path = 'save_path' in step_params and step_params.get('save_path')
                    
                    if has_save_to_desktop and not has_save_path:
                        logger.warning(f"⚠️ 步骤 {i} screenshot_desktop：用户要求保存到桌面，但未传递save_path参数，自动添加")
                        step_params['save_path'] = "~/Desktop/screenshot.png"
                        steps[i-1]['params'] = step_params
                        logger.info("✅ 已自动添加 save_path: ~/Desktop/screenshot.png")
            
            # 后处理3：检查 browser_screenshot 是否需要添加 save_path
            for i, step in enumerate(steps, 1):
                step_type = step.get('type')
                step_params = step.get('params', {})
                
                # 如果是 browser_screenshot，检查用户是否要求保存到桌面
                if step_type == 'browser_screenshot':
                    # 检查用户指令中是否包含"保存到桌面"、"保存桌面"等关键词
                    has_save_to_desktop = (
                        "保存到桌面" in user_instruction or
                        "保存桌面" in user_instruction or
                        "保存到 ~/Desktop" in user_instruction or
                        "save to desktop" in user_instruction_lower or
                        "save desktop" in user_instruction_lower or
                        ("保存" in user_instruction and "桌面" in user_instruction) or
                        ("save" in user_instruction_lower and "desktop" in user_instruction_lower)
                    )
                    
                    # 检查是否已经传递了 save_path 参数
                    has_save_path = 'save_path' in step_params and step_params.get('save_path')
                    
                    if has_save_to_desktop and not has_save_path:
                        logger.warning(f"⚠️ 步骤 {i} browser_screenshot：用户要求保存到桌面，但未传递save_path参数，自动添加")
                        step_params['save_path'] = "~/Desktop/screenshot.png"
                        steps[i-1]['params'] = step_params
                        logger.info("✅ 已自动添加 save_path: ~/Desktop/screenshot.png")
            
            return steps
            
        except anthropic.APIError as e:
            logger.error(f"Claude API调用失败: {e}")
            raise PlannerError(f"规划失败: Claude API错误 - {e}")
        except Exception as e:
            logger.error(f"规划失败: {e}", exc_info=True)
            raise PlannerError(f"规划失败: {e}")
    
    def _build_prompt(
        self,
        instruction: str,
        context: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        构建规划提示词 - 简化版
        
        核心理念：
        - 简洁清晰，不硬编码规则
        - 告诉 AI 它有什么能力
        - 让 AI 自己决定怎么做
        
        Args:
            instruction: 用户指令
            context: 上下文信息
        
        Returns:
            提示词
        """
        # 获取用户主目录
        home_dir = str(Path.home())
        
        # 构建上下文信息（增强版：帮助AI理解指代词）
        context_str = ""
        if context:
            context_parts = []
            
            # 0. 处理记忆上下文（AI对用户的长期记忆）
            memory_context = context.get("memory_context", "")
            if memory_context:
                context_parts.append(f"""### AI记忆（用户偏好和习惯）
{memory_context}

**重要**: 请根据这些记忆信息，更好地理解用户的习惯和偏好，生成符合用户期望的执行计划。""")
            
            # 0.5 处理工作流建议
            workflow_suggestion = context.get("workflow_suggestion")
            if workflow_suggestion:
                pattern = workflow_suggestion.get("pattern", {})
                action_seq = pattern.get("action_sequence", [])
                if action_seq:
                    context_parts.append(f"""### 工作流建议
用户经常执行类似任务，建议参考之前成功的步骤模式：
步骤序列：{' → '.join(action_seq)}
成功率：{pattern.get('success_rate', 0) * 100:.0f}%""")
            
            # 1. 处理最近创建/操作的文件
            last_file = context.get("last_created_file")
            created_files = context.get("created_files", [])
            
            if last_file:
                context_parts.append(f"""### 最近操作的文件
**路径**: `{last_file}`
**重要**: 当用户说"这个文件"、"那个文件"、"刚才的文件"、"这张图片"、"那张截图"等指代词时，
指的就是这个文件：`{last_file}`""")
            
            if created_files and len(created_files) > 1:
                files_list = "\n".join([f"- `{f}`" for f in created_files[-5:]])  # 只显示最近5个
                context_parts.append(f"""### 本次会话创建的所有文件
{files_list}""")
            
            # 2. 处理用户附加的路径
            attached_path = context.get("attached_path")
            if attached_path:
                context_parts.append(f"""### 用户指定的目标路径
用户拖拽或选择了这个路径：`{attached_path}`
请在任务中使用这个路径。""")
            
            # 3. 处理聊天历史（让AI理解对话上下文）
            chat_history = context.get("chat_history", [])
            if chat_history:
                history_parts = []
                for msg in chat_history[-5:]:  # 只保留最近5条
                    role = "用户" if msg.get("role") == "user" else "AI"
                    content = msg.get("content", "")
                    # 截断过长的内容
                    if len(content) > 200:
                        content = content[:200] + "..."
                    history_parts.append(f"**{role}**: {content}")
                
                if history_parts:
                    context_parts.append(f"""### 对话历史
{chr(10).join(history_parts)}

**重要**: 请根据对话历史理解用户的意图。如果用户提到"这个"、"那个"、"它"等指代词，
请结合对话历史和上面的"最近操作的文件"来确定具体指的是什么。""")
            
            if context_parts:
                context_str = "\n## 上下文信息\n\n" + "\n\n".join(context_parts) + "\n"
        
        # 按需精简 prompt：根据任务类型只包含相关工具说明
        needs_browser = any(kw in instruction for kw in [
            "网页", "网站", "浏览", "搜索", "下载", "http", "www",
            "百度", "谷歌", "google", "访问", "登录",
        ])
        needs_word = any(kw in instruction.lower() for kw in [
            "word", "docx", ".docx", "替换文字", "替换文档",
        ])
        needs_chart = any(kw in instruction for kw in [
            "图表", "柱形图", "饼图", "折线图", "统计", "chart",
        ])
        
        browser_section = ""
        if needs_browser:
            browser_section = """
### 1. 浏览器操作
- browser_navigate: 访问网页，params: {{url: "..."}}
- browser_click: 点击元素，params: {{selector: "..."}} 或 {{text: "..."}}
- browser_input: 输入文本，params: {{selector: "...", text: "..."}}
- browser_screenshot: 截取当前浏览器页面，params: {{save_path: "保存路径（可选）"}}
- download_file: 下载文件（通过浏览器点击下载链接），params: {{selector: "..."}} 或 {{text: "..."}}, 可选 {{save_path: "保存路径/目录"}}, {{timeout: 60000}}
- request_login: 请求用户登录（弹出登录对话框），params: {{site_name: "网站名", username_selector: "用户名输入框选择器", password_selector: "密码输入框选择器", submit_selector: "提交按钮选择器（可选）"}}
- request_captcha: 请求验证码（截取验证码图片并弹出输入框），params: {{site_name: "网站名", captcha_image_selector: "验证码图片选择器", captcha_input_selector: "验证码输入框选择器"}}

**重要规则：截图工具选择**
- **browser_screenshot**：用于截图浏览器中的网页内容
  - 当用户指令包含"搜索"、"打开网页"、"访问"、"浏览"等浏览器操作，然后说"截图"时，必须使用 browser_screenshot
  - 示例：用户说"搜索GitHub然后截图" → 使用 browser_screenshot（截图GitHub网页）
  - 示例：用户说"打开百度截图给我" → 使用 browser_screenshot（截图百度网页）
- **screenshot_desktop**：用于截图整个桌面（不是浏览器页面）
  - 只有当用户明确说"截图桌面"、"截图整个屏幕"、"桌面截图"时才使用
  - 示例：用户说"截图桌面" → 使用 screenshot_desktop
  - 示例：用户说"截图整个屏幕" → 使用 screenshot_desktop
- **关键区别**：
  - "保存到桌面" ≠ "截图桌面"！"保存到桌面"只是指保存路径，不是截图对象
  - 如果用户先有浏览器操作（如"搜索GitHub"），然后说"截图给我，保存到桌面"，应该用 browser_screenshot + save_path参数，而不是 screenshot_desktop
"""
        
        word_doc_section = ""
        if needs_word:
            word_doc_section = """
**脚本示例（Word文档文字替换）**：
**警告：.docx 是 ZIP 压缩包，绝对禁止用 open() 读取！必须用 python-docx 库！**
**重要：Word 可能把一个名字拆成多个 runs（格式块）！仅在 run.text 内 replace 仍可能"替换 0 处"。下面是支持跨 runs 的完美替换写法。**
    ```python
import json
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(json.dumps({{"success": False, "message": "请先安装 python-docx: pip install python-docx"}}))
    exit(0)

def replace_across_runs(paragraph, old_text, new_text):
    runs = paragraph.runs
    if not runs:
        return 0
    replaced = 0
    while True:
        full = "".join([r.text for r in runs])
        idx = full.find(old_text)
        if idx == -1:
            break
        mapping = []
        for run_i, r in enumerate(runs):
            for off in range(len(r.text)):
                mapping.append((run_i, off))
        start = idx
        end = idx + len(old_text) - 1
        if end >= len(mapping):
            break
        s_run, s_off = mapping[start]
        e_run, e_off = mapping[end]
        before = runs[s_run].text[:s_off]
        after = runs[e_run].text[e_off + 1:]
        if s_run == e_run:
            runs[s_run].text = before + new_text + after
        else:
            runs[s_run].text = before + new_text
            for j in range(s_run + 1, e_run):
                runs[j].text = ""
            runs[e_run].text = after
        replaced += 1
    return replaced

file_path = Path.home() / "Desktop" / "文档.docx"
old_text = "原文字"
new_text = "新文字"
doc = Document(file_path)
count = 0
for para in doc.paragraphs:
    if old_text in para.text:
        count += replace_across_runs(para, old_text, new_text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        count += replace_across_runs(para, old_text, new_text)
for section in doc.sections:
    for para in section.header.paragraphs:
        if old_text in para.text:
            count += replace_across_runs(para, old_text, new_text)
    for para in section.footer.paragraphs:
        if old_text in para.text:
            count += replace_across_runs(para, old_text, new_text)
doc.save(file_path)
if count == 0:
    result = {{"success": False, "message": "替换完成但未找到任何可替换内容（0 处）"}}
else:
    result = {{"success": True, "message": "替换完成，共替换 " + str(count) + " 处"}}
print(json.dumps(result, ensure_ascii=False))
```
"""
        
        chart_section = ""
        if needs_chart:
            chart_section = """
## Matplotlib 图表绑定用法（画图必看）

1. 饼图颜色使用列表，不要用 cm.Set3：
   ```python
   colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE']
   plt.pie(sizes, labels=labels, colors=colors[:len(labels)], autopct='%1.1f%%')
   ```
2. 中文显示：`plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'SimHei']`
3. 保存图表：`plt.savefig(路径, dpi=150, bbox_inches='tight', facecolor='white')`
4. **不要使用** `plt.cm.set3`、`plt.cm.Set3` 等，会报错！
"""
        
        prompt = f"""你是 DeskJarvis，一个智能桌面助手。请用中文思考和输出。

**核心原则**：
- **理解用户的真实意图**：仔细分析用户的自然语言指令，理解用户想做什么
- **拆分多个操作**：如果用户指令包含多个操作（如"打开应用然后输入文本"），必须拆分为多个步骤
- **每个步骤只做一件事**：一个步骤只执行一个操作

## 你的能力
{browser_section}
### 2. 系统操作
- screenshot_desktop: 截取整个桌面（不是浏览器页面），params: {{save_path: "保存路径（可选）"}}
  - **注意**：只有当用户明确要求"截图桌面"、"截图整个屏幕"时才使用此工具
  - 如果用户先有浏览器操作（如"搜索"、"打开网页"），然后说"截图"，应该使用 browser_screenshot，而不是 screenshot_desktop
- open_folder: 打开文件夹，params: {{folder_path: "..."}}
- open_file: 打开文件，params: {{file_path: "..."}}
- open_app: 打开应用，params: {{app_name: "应用名称"}}
- close_app: 关闭应用，params: {{app_name: "应用名称"}}

### 2.5 文件操作工具（简单操作可用，复杂操作建议用脚本）
- file_read: 读取文件，params: {{file_path: "文件路径"}}
- file_write: 写入文件，params: {{file_path: "文件路径", content: "内容"}}
- file_create: 创建文件，params: {{file_path: "文件路径", content: "内容"}}
- file_rename: 重命名文件，params: {{file_path: "原文件路径", new_name: "新文件名"}}
- file_move: 移动文件，params: {{file_path: "原路径", destination: "目标路径"}}
- file_copy: 复制文件，params: {{file_path: "原路径", destination: "目标路径"}}
- file_delete: 删除文件，params: {{file_path: "文件路径"}}

### 2.6 系统控制工具
- set_volume: 设置音量，params: {{level: 0-100}} 或 {{action: "mute/unmute/up/down"}}
- set_brightness: 设置屏幕亮度，params: {{level: 0.0-1.0}} 或 {{action: "up/down/max/min"}}
- send_notification: 发送系统通知，params: {{title: "标题", message: "内容"}}
- speak: 语音播报，params: {{text: "要播报的内容"}}
- clipboard_read: 读取剪贴板，params: {{}}
- clipboard_write: 写入剪贴板，params: {{content: "内容"}}
- keyboard_type: 键盘输入，params: {{text: "要输入的文本"}}
- keyboard_shortcut: 按键/快捷键（用于回车/Tab/Esc/方向键/⌘C 等），params: {{keys: "command+c"}}, 可选 {{repeat: 2}}（如按两次回车）

**键盘规则（重要！）**：
- **输入文字**用 `keyboard_type`（支持中文、英文、数字、符号）
  - 示例：输入"张旭政" → `{{"type":"keyboard_type","params":{{"text":"张旭政"}}}}`
  - 示例：输入"zhangxuzheng" → `{{"type":"keyboard_type","params":{{"text":"zhangxuzheng"}}}}`
- **按回车/Tab/Esc/方向键**必须用 `keyboard_shortcut`，不要把 "enter" 当文本输入！
  - 按两次回车：`{{"type":"keyboard_shortcut","params":{{"keys":"enter","repeat":2}}}}`
- mouse_click: 鼠标点击，params: {{x: 100, y: 200}}
- window_minimize: 最小化窗口，params: {{app_name: "应用名（可选）"}}
- window_maximize: 最大化窗口，params: {{app_name: "应用名（可选）"}}

**下载工具（推荐，避免脚本语法错误）**：
- download_latest_python_installer: 下载最新 Python 安装包，params: {{save_dir: "保存目录（可选，默认桌面）"}} 或 {{save_path: "保存路径/目录（可选）"}}, 可选 {{timeout: 180000}}

**系统信息和图片处理**：
- get_system_info: 获取系统信息，params: {{info_type: "battery/disk/memory/apps/network/all", save_path: "~/Desktop/系统报告.md（可选，指定后自动保存）"}}
  **重要：查询系统信息必须使用这个工具，不要自己写脚本！如果用户要求保存，直接在 save_path 中指定路径！**
- image_process: 图片处理，params: {{image_path: "图片路径", action: "compress/resize/convert/info", width: 800, height: 600, format: "jpg/png/webp", quality: 80}}

**定时提醒**：
- set_reminder: 设置提醒，params: {{message: "提醒内容", delay: "5分钟/1小时/30秒", repeat: "daily/hourly（可选）"}}
- list_reminders: 列出提醒，params: {{}}
- cancel_reminder: 取消提醒，params: {{reminder_id: "提醒ID"}}

**工作流管理**：
- create_workflow: 创建工作流，params: {{name: "工作流名", commands: ["命令1", "命令2"], description: "描述"}}
- list_workflows: 列出工作流，params: {{}}
- delete_workflow: 删除工作流，params: {{name: "工作流名"}}

**任务历史**：
- get_task_history: 获取历史，params: {{limit: 20}}
- search_history: 搜索历史，params: {{keyword: "关键词"}}
- add_favorite: 添加收藏，params: {{instruction: "指令内容", name: "收藏名（可选）"}}
- list_favorites: 列出收藏，params: {{}}
- remove_favorite: 移除收藏，params: {{favorite_id: "收藏ID"}}

**文本AI处理**：
- text_process: AI文本处理，params: {{text: "要处理的文本", action: "translate/summarize/polish/expand/fix_grammar", target_lang: "目标语言（翻译时使用）"}}

### 3. Python 脚本（推荐用于文件操作和批量操作）
使用 execute_python_script 执行自定义 Python 代码。

execute_python_script 的 params:
- script: base64 编码的 Python 脚本（必须）
- reason: 为什么需要脚本（必须）
- safety: 安全说明（必须）

**脚本必须遵循的规范**：
1. 必须是完整可执行的 Python 代码
2. 输出格式必须是：print(json.dumps({{"success": True, "message": "xxx"}})) 或 print(json.dumps({{"success": False, "message": "xxx"}}))
3. **本地文件统计/生成图表必须用 execute_python_script**，不要使用任何 browser_* 工具
3. **必须通过 ruff 快检（E/F/B）**：系统会在执行前自动运行 `ruff check --select E,F,B`，不通过会直接失败并进入反思重试
   - 常见必修点：只 import 你真正用到的（避免 F401），不要引用未定义变量（F821），不要 `except:`（E722），确保没有语法错误（E999），`raise` 保留异常链（B904）
4. **禁止使用 f-string**！所有字符串拼接必须用 + 号
5. **语法检查（极其重要）**：
   - 每个引号必须配对闭合
   - 每个括号必须配对闭合
   - **每个 try 必须有 except**（最常见错误！）
   - 字符串拼接格式: "文字" + str(变量) + "文字"
   - **平台检测**：`import sys; sys.platform == "darwin"/"win32"/"linux"`
6. 注释使用中文，变量名使用英文
7. 用户主目录: {home_dir}
8. **错误处理**：所有操作必须在 try-except 中，except 中输出 JSON 格式的错误信息
9. **HTTP 请求（重要！）**：
   - **必须使用 requests 库**，不要用 urllib！urllib 不会自动解压 gzip！
   - `import requests` → `response = requests.get(url)`
   - 下载二进制文件用 `response.content`，下载文本用 `response.text`
   - 示例：
     ```python
     import requests
     response = requests.get(url)
     html = response.text  # 文本
     with open(path, "wb") as f:
         f.write(response.content)  # 二进制文件
     ```

**脚本示例（删除文件）**：
```python
import os
import json
import glob
from pathlib import Path

# 目标目录
desktop = Path.home() / "Desktop"
pattern = str(desktop / "*screenshot*.png")

# 查找匹配文件
files = glob.glob(pattern)
deleted = []
errors = []

for f in files:
    try:
        os.remove(f)
        deleted.append(os.path.basename(f))
    except Exception as e:
        errors.append(str(e))

# 输出结果（必须是有效JSON）
if deleted:
    result = {{"success": True, "message": "成功删除 " + str(len(deleted)) + " 个文件", "data": {{"deleted": deleted}}}}
else:
    result = {{"success": True, "message": "没有找到匹配的文件", "data": {{}}}}

print(json.dumps(result, ensure_ascii=False))
```

**脚本示例（重命名文件）**：
```python
import os
import json
from pathlib import Path

# 目标目录
folder = Path.home() / "Desktop" / "images"

# 获取所有图片并按大小排序
images = [(f, f.stat().st_size) for f in folder.iterdir() if f.suffix.lower() in [".png", ".jpg", ".jpeg"]]
images.sort(key=lambda x: x[1])

# 重命名
renamed = []
for i, (img, size) in enumerate(images, 1):
    new_name = folder / (str(i) + img.suffix)
    os.rename(img, new_name)
    renamed.append({{"old": img.name, "new": new_name.name}})

result = {{"success": True, "message": "成功重命名 " + str(len(renamed)) + " 个文件", "data": {{"renamed": renamed}}}}
print(json.dumps(result, ensure_ascii=False))
```

**重要：message 字符串必须使用字符串拼接（"xxx " + str(n) + " yyy"），不要使用 f-string！**

**脚本示例（搜索并操作文件 - 当用户说的文件名可能不完整时）**：
      ```python
      import os
import json
from pathlib import Path

# 用户说"强制执行申请书"，可能是部分文件名，需要搜索
desktop = Path.home() / "Desktop"
keyword = "强制执行"  # 从用户指令中提取关键词

# 搜索匹配的文件
matches = []
for f in desktop.iterdir():
    if f.is_file() and keyword in f.name:
        matches.append(f)

if not matches:
    result = {{"success": False, "message": "未找到包含'" + keyword + "'的文件"}}
elif len(matches) > 1:
    files_str = ", ".join([m.name for m in matches])
    result = {{"success": False, "message": "找到多个匹配文件: " + files_str + "，请指定具体文件"}}
else:
    target = matches[0]
    # 然后对 target 进行操作...
    result = {{"success": True, "message": "找到文件: " + target.name}}

print(json.dumps(result, ensure_ascii=False))
```

{word_doc_section}
## 任务
{instruction}
{context_str}

## 文件名理解规则

当用户提到文件名时：
1. **用户说的可能是部分文件名**（如"强制执行申请书"可能指"强制执行申请书-张三.docx"）
2. **先搜索匹配的文件**，再进行操作
3. **不要猜测完整文件名**，使用关键词搜索

## 关键警告

1. **绝对禁止 f-string**：不要用 f"xxx{{变量}}xxx" 格式！会导致编码错误！必须用 "xxx" + str(变量) + "xxx"
2. **JSON 输出必须正确**：脚本最后一行必须是 print(json.dumps(..., ensure_ascii=False))
3. **错误也要 JSON 格式**：except 中也要用 print(json.dumps({{"success": False, "message": str(e)}})) 输出

{chart_section}
## 输出格式

返回一个 JSON 数组，每个元素是一个步骤：
```json
[
  {{
    "type": "步骤类型（如 open_app、keyboard_type、keyboard_shortcut、execute_python_script、screenshot_desktop、open_folder 等）",
    "action": "简短操作描述（中文，如：打开企业微信）",
    "params": {{}},
    "description": "给用户看的详细描述（中文）"
  }}
]
```

**重要提示**：
- 如果用户说"打开XXX然后YYY"或"打开XXX YYY"，XXX是应用名，YYY是后续操作，必须拆分为多个步骤
- 例如："打开企业微信控制键盘输入zhangxuzheng按空格" → 应该拆分为3个步骤：
  1. open_app（app_name: "企业微信"）
  2. keyboard_type（text: "zhangxuzheng"）
  3. keyboard_shortcut（keys: "space"）

## 重要规则

1. **所有文件操作使用 execute_python_script**：包括重命名、删除、移动、复制、批量处理等
2. **脚本代码必须正确**：检查变量名、语法、缩进，确保可执行
3. **中文描述**：所有 description 和 action 使用中文
4. **路径处理**：使用 Path 对象处理路径，支持 ~ 和中文路径
5. **只返回 JSON 数组**：不要有任何其他解释文字
6. **理解自然语言**：仔细分析用户指令，正确拆分多个操作

现在请规划任务："""
        
        return prompt
    
    def _call_reflection_api(self, prompt: str) -> Dict[str, Any]:
        """
        调用Claude API进行反思
        
        Args:
            prompt: 反思提示词
        
        Returns:
            包含分析和新计划的字典
        """
        logger.info("调用Claude进行反思...")
        
        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}]
            )
            
            content = response.content[0].text
            logger.debug(f"反思响应: {content[:500]}...")
            
            # 解析反思结果
            return self._parse_reflection_response(content)
            
        except Exception as e:
            logger.error(f"反思API调用失败: {e}")
            raise PlannerError(f"反思失败: {e}")
    
    def _parse_reflection_response(self, content: str) -> Dict[str, Any]:
        """
        解析反思响应
        
        Args:
            content: AI响应内容
        
        Returns:
            解析后的反思结果
        """
        content = content.strip()
        
        # 移除markdown代码块
        if content.startswith("```"):
            lines = content.split("\n")
            if len(lines) > 2:
                content = "\n".join(lines[1:-1])
        
        # 尝试提取JSON对象
        start_idx = content.find('{')
        end_idx = content.rfind('}')
        
        if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
            content = content[start_idx:end_idx + 1]
        
        try:
            result = json.loads(content)
            
            # 验证格式
            if "analysis" not in result:
                result["analysis"] = "无分析"
            if "new_plan" not in result:
                result["new_plan"] = []
            
            # 验证 new_plan 是列表
            if not isinstance(result["new_plan"], list):
                result["new_plan"] = []
            
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"解析反思响应失败: {e}")
            return {
                "analysis": f"解析失败: {e}",
                "new_plan": []
            }
